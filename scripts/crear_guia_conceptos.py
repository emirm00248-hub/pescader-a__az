from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "Guia_Conceptos_POO_MVC_SOLID_Heap.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(46, 116, 181) if level <= 2 else RGBColor(31, 77, 120)
    return p


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(31, 58, 95)
    p.add_run("\n" + body).font.size = Pt(9)
    doc.add_paragraph()


def add_concept_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_widths(table, [1.25, 1.45, 1.7, 2.1])
    headers = ["Concepto", "Donde esta", "Archivo / lineas", "Como se demuestra"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color="FFFFFF")
        set_cell_shading(table.rows[0].cells[i], "1F4D78")
    for concept, where, lines, proof in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], concept, bold=True)
        set_cell_text(cells[1], where)
        set_cell_text(cells[2], lines)
        set_cell_text(cells[3], proof)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.15

    title = doc.add_paragraph()
    title_run = title.add_run("Guia de conceptos en el codigo")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(11, 37, 69)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Proyecto: Pescaderia Azcorra - ASP.NET Core")
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(85, 85, 85)

    add_note(
        doc,
        "Idea principal",
        "Este documento muestra donde se encuentran los conceptos de POO, MVC, SOLID, patrones y memoria en tu codigo. "
        "La intencion es que puedas explicarlos con evidencia real: archivo, linea y comportamiento.",
    )

    add_heading(doc, "Resumen rapido", 1)
    add_bullets(
        doc,
        [
            "La POO se ve principalmente en Models, Services y Patterns.",
            "MVC aparece como una version web API: Models y Controllers estan claros; la vista esta como frontend estatico en wwwroot, no como Razor Views.",
            "SOLID aparece sobre todo por separacion de responsabilidades y por depender de IOrderService en los controladores.",
            "Heap y memoria aparecen conceptualmente cuando se crean objetos con new, listas en memoria y el Singleton InMemoryDatabase.",
            "Los patrones usados son Factory, Singleton y Observer.",
        ],
    )

    add_heading(doc, "POO: Programacion Orientada a Objetos", 1)
    add_concept_table(
        doc,
        [
            (
                "Clase",
                "Modelos y servicios",
                "Models/Product.cs:3; Models/Order.cs:19; Services/OrderService.cs:8",
                "El sistema esta dividido en moldes de objetos: Product, Order, OrderService, etc.",
            ),
            (
                "Objeto",
                "Creacion de productos, ordenes y servicios",
                "Patterns/Singleton/InMemoryDatabase.cs:17-18; Services/OrderService.cs:67-78",
                "Se crean instancias reales en memoria: listas de productos, diccionario de ordenes y nuevas ordenes.",
            ),
            (
                "Abstraccion",
                "Clase base e interfaz",
                "Models/Product.cs:3; Services/IOrderService.cs:6",
                "Product define lo comun de todos los productos; IOrderService define lo que el servicio debe hacer sin amarrarse a una implementacion concreta.",
            ),
            (
                "Encapsulamiento",
                "Propiedades protegidas o privadas",
                "Models/Product.cs:5-18; Models/Order.cs:21-32",
                "Los datos no quedan como campos libres; se controlan con protected set o private set para evitar cambios directos desde fuera.",
            ),
            (
                "Herencia",
                "Productos especializados",
                "Models/DishProduct.cs:3; Models/SeafoodProduct.cs:3; Models/MiscProduct.cs:3",
                "DishProduct, SeafoodProduct y MiscProduct heredan de Product.",
            ),
            (
                "Polimorfismo",
                "Metodos virtuales y override",
                "Models/Product.cs:40,45; Models/DishProduct.cs:12; Models/SeafoodProduct.cs:12; Models/MiscProduct.cs:12",
                "Cada tipo de producto puede responder de forma distinta a GetProductSummary; GetCurrentPrice se llama desde Product sin depender de una clase hija concreta.",
            ),
            (
                "Constructor",
                "Inicializacion de objetos",
                "Models/Product.cs:20-37; Models/Order.cs:34-48; Services/OrderService.cs:13",
                "Los constructores obligan a crear objetos con la informacion necesaria desde el inicio.",
            ),
        ],
    )

    add_heading(doc, "MVC en este proyecto", 1)
    add_note(
        doc,
        "Aclaracion importante",
        "El proyecto no usa MVC clasico completo con Views Razor. Usa ASP.NET Core Web API: los Controllers atienden rutas, los Models representan datos, "
        "y la vista visual esta en wwwroot como HTML, CSS y JavaScript.",
    )
    add_concept_table(
        doc,
        [
            (
                "Model",
                "Clases de dominio y DTOs",
                "Models/Product.cs; Models/Order.cs; DTOs/OrderInputs.cs",
                "Representan productos, pedidos y datos que entran desde el formulario.",
            ),
            (
                "View",
                "Frontend estatico",
                "wwwroot/index.html; wwwroot/style.css; wwwroot/app.js",
                "La interfaz que ve el usuario esta hecha con HTML, CSS y JS.",
            ),
            (
                "Controller",
                "Controladores de API",
                "Controllers/ProductsController.cs:6-27; Controllers/OrdersController.cs:7-23",
                "Reciben peticiones HTTP, llaman al servicio y devuelven respuestas.",
            ),
            (
                "Rutas",
                "Atributos de ASP.NET Core",
                "ProductsController.cs:7,17,24; OrdersController.cs:8,18",
                "Define URLs como api/products y api/orders.",
            ),
        ],
    )

    add_heading(doc, "SOLID", 1)
    add_concept_table(
        doc,
        [
            (
                "SRP",
                "Responsabilidad unica",
                "Controllers/*.cs; Services/OrderService.cs:31-82; Models/*.cs",
                "Los controladores manejan HTTP, el servicio maneja reglas de negocio y los modelos representan datos.",
            ),
            (
                "OCP",
                "Abierto a extension",
                "Patterns/Observer/IOrderObserver.cs:5; Patterns/Observer/OrderSubject.cs:9-21",
                "Se pueden agregar nuevos observadores sin reescribir el flujo principal de crear ordenes.",
            ),
            (
                "LSP",
                "Sustitucion por herencia",
                "Models/DishProduct.cs:3; Models/SeafoodProduct.cs:3; Models/MiscProduct.cs:3",
                "Las clases hijas pueden usarse donde se espera un Product.",
            ),
            (
                "ISP",
                "Interfaz pequena",
                "Services/IOrderService.cs:6-10",
                "La interfaz solo expone operaciones necesarias: obtener productos y crear ordenes.",
            ),
            (
                "DIP",
                "Depender de abstracciones",
                "Program.cs:15; ProductsController.cs:10-14; OrdersController.cs:11-15",
                "Los controladores reciben IOrderService, no crean directamente OrderService.",
            ),
        ],
    )

    add_heading(doc, "Patrones de diseno", 1)
    add_concept_table(
        doc,
        [
            (
                "Factory",
                "Creacion centralizada de productos",
                "Patterns/Factory/ProductFactory.cs:5-24; InMemoryDatabase.cs:44-63",
                "ProductFactory decide si crea DishProduct, SeafoodProduct o MiscProduct segun la categoria.",
            ),
            (
                "Singleton",
                "Base de datos en memoria unica",
                "Patterns/Singleton/InMemoryDatabase.cs:7-32",
                "La propiedad Instance controla que exista una unica instancia compartida.",
            ),
            (
                "Observer",
                "Notificacion al crear orden",
                "Patterns/Observer/OrderSubject.cs:5-21; ConsoleLogObserver.cs:5-10; OrderService.cs:17-18,82",
                "Cuando se crea una orden, OrderSubject notifica a los observadores registrados.",
            ),
        ],
    )

    add_heading(doc, "Memoria: heap, stack y garbage collector", 1)
    add_concept_table(
        doc,
        [
            (
                "Heap",
                "Objetos creados con new",
                "OrderService.cs:17,54-59; InMemoryDatabase.cs:17-18,30",
                "Los objetos como OrderSubject, List<Product>, ConcurrentDictionary y productos viven en memoria administrada.",
            ),
            (
                "Stack",
                "Variables locales y parametros",
                "OrderService.cs:31-82",
                "Variables como subtotal, tax, total, itemInput y product son referencias o valores locales del metodo CreateOrder.",
            ),
            (
                "Garbage Collector",
                "Liberacion automatica",
                "Patterns/Singleton/InMemoryDatabase.cs:37-40",
                "C# administra memoria automaticamente; el finalizador muestra que el objeto podria limpiarse cuando el GC lo recolecte.",
            ),
            (
                "Colecciones en memoria",
                "List y ConcurrentDictionary",
                "InMemoryDatabase.cs:12-18; OrderService.cs:80",
                "Products guarda el catalogo y Orders guarda pedidos en memoria durante la ejecucion.",
            ),
        ],
    )

    add_heading(doc, "API y flujo del pedido", 1)
    add_concept_table(
        doc,
        [
            (
                "Inyeccion de dependencias",
                "Registro y uso del servicio",
                "Program.cs:15; Controllers/*.cs",
                "ASP.NET crea el servicio y lo entrega a los controladores.",
            ),
            (
                "DTO",
                "Datos de entrada del pedido",
                "DTOs/OrderInputs.cs:3-17",
                "OrderInput y OrderItemInput transportan lo que manda el frontend al backend.",
            ),
            (
                "Validacion",
                "Reglas antes de crear pedido",
                "Services/OrderService.cs:31-49",
                "Se valida que el pedido tenga cliente, items y productos existentes.",
            ),
            (
                "Calculo de precio",
                "Uso de metodo del modelo",
                "Services/OrderService.cs:51-52; Models/Product.cs:40-43",
                "El servicio pide al producto su precio actual, incluyendo oferta si existe.",
            ),
            (
                "Persistencia temporal",
                "Guardar pedido en memoria",
                "Services/OrderService.cs:80; InMemoryDatabase.cs:13",
                "El pedido se guarda en un ConcurrentDictionary mientras la app esta ejecutandose.",
            ),
        ],
    )

    add_heading(doc, "Que puedes decir en la exposicion", 1)
    add_bullets(
        doc,
        [
            "Mi proyecto separa la interfaz, los controladores, los modelos y la logica de negocio.",
            "La POO no esta solo en comentarios: se ve en Product como clase abstracta, en las clases hijas y en los metodos override.",
            "Uso una interfaz IOrderService para que los controladores dependan de una abstraccion.",
            "Uso Factory para crear productos segun su categoria, Singleton para una base en memoria y Observer para notificar cuando se crea una orden.",
            "Los objetos creados con new viven en memoria administrada; C# usa garbage collector para limpiar objetos que ya no se usan.",
        ],
    )

    add_heading(doc, "Observaciones honestas", 1)
    add_bullets(
        doc,
        [
            "MVC aqui es una adaptacion con Web API y frontend estatico, no MVC tradicional con Razor Views.",
            "La base de datos es en memoria, por eso los pedidos se pierden al reiniciar la app.",
            "GetProductSummary demuestra polimorfismo, aunque todavia no se muestra en el frontend.",
            "El Observer actual solo escribe en consola; podria crecer con email, SMS o registro externo.",
        ],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Checklist rapido de archivos", 1)
    add_concept_table(
        doc,
        [
            ("Program.cs", "Configuracion", "Program.cs:13-24", "Registra controladores, servicio, archivos estaticos y rutas."),
            ("Controllers", "MVC Controller", "Controllers/*.cs", "Reciben peticiones HTTP y llaman al servicio."),
            ("DTOs", "Entrada de datos", "DTOs/OrderInputs.cs:3-17", "Define la forma del pedido que llega desde JS."),
            ("Models", "POO principal", "Models/*.cs", "Clases, herencia, encapsulamiento y polimorfismo."),
            ("Services", "Logica de negocio", "Services/IOrderService.cs; Services/OrderService.cs", "Reglas para productos y pedidos."),
            ("Patterns", "Patrones", "Patterns/Factory; Patterns/Singleton; Patterns/Observer", "Factory, Singleton y Observer."),
            ("wwwroot", "Vista/frontend", "wwwroot/index.html; app.js; style.css", "Interfaz visual, carrito y llamadas a la API."),
        ],
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
